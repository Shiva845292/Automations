import time
from _msi import OpenDatabase, MSIDBOPEN_READONLY
import os
import docx
from datetime import date
import aspose.words as aw
from docx.shared import Pt
from docx.enum.text import WD_COLOR_INDEX


"""
import math
size = 0
Folderpath = "\\"+"\\"+Location+"\\"+ UPN
for path, dirs, files in os.walk(Folderpath):
    for f in files:
        fp = os.path.join(path, f)
        size += os.stat(fp).st_size
        if size == 0:
            print("0B")
        else:
            size_name = ("B", "KB", "MB", "GB", "TB", "PB", "EB", "ZB", "YB")
            i = int(math.floor(math.log(size, 1024)))
            p = math.pow(1024, i)
            s = round(size / p, 2)
            ss = ("%s %s" % (s, size_name[i]))
print(ss)
"""

today = date.today()
today = today.strftime("%d/%m/%Y")
application = []
for file in os.listdir():
    if (file.startswith("ZIC-") or file.startswith("ZCX-")) and (file.endswith(".docx")):
        doc = docx.Document(file)
        for index, table in enumerate(doc.tables):
            if index == 0:
                for row in range(len(table.rows)):
                    for col in range(len(table.columns)):
                        application.append(table.cell(row, col).text)
            elif index == 11:
                for row in range(len(table.rows)):
                    for col in range(len(table.columns)):
                        Ugrades = table.cell(row, col).text
            elif index == 14:
                for row in range(len(table.rows)):
                    for col in range(len(table.columns)):
                        pre = table.cell(row, col).text
            elif index == 16:
                for row in range(len(table.rows)):
                    for col in range(len(table.columns)):
                        comp = table.cell(row, col).text
        print("Read install Guide is done......!")
        eACS = file
        n = len(eACS)
        doc = aw.Document(str(eACS[0:n - 18]) + "_eACS.rtf")
        doc.save("Dummy_eACS.docx")
        rtf = []
        for file in os.listdir():
            if file.endswith(".docx") and ("_eACS" in file):
                doc = docx.Document(file)
                for index, table in enumerate(doc.tables):
                    for row in range(len(table.rows)):
                        for col in range(len(table.columns)):
                            rtf.append(table.cell(row, col).text)
                os.remove("Dummy_eACS.docx")
        print("reading eACS.....!")

l = []
for file in os.listdir():
    if (file.startswith("ZIC-") or file.startswith("ZCX-")) and (file.endswith(".exe")):
        l.append(file)
        break
    elif (file.startswith("ZIC-") or file.startswith("ZCX-")) and (file.endswith(".msi")):
        l.append(file)
        break
    elif (file.startswith("ZIC-") or file.startswith("ZCX-")) and (file.endswith(".mst")):
        l.append(file)
        break
    else:
        pass
if len(l) == 1 and l[0].endswith(".msi"):
    MSI = l[0]
    if MSI[:4] == "ZIC-":
        Location = "172.16.64.2\Pkg-I\Zurich\pkg"
    else:
        Location = "172.16.64.2\Pkg-I\Zurich\pkg-citrix"

    if MSI.endswith(".msi"):
        def GetMsiProperty(path, property):
            db = OpenDatabase(path, MSIDBOPEN_READONLY)
            view = db.OpenView("SELECT Value FROM Property WHERE Property='" + property + "'")
            view.Execute(None)
            result = view.Fetch()
            return result.GetString(1)


        def Summaryinfo(path, n):
            dbobject = OpenDatabase(path, MSIDBOPEN_READONLY)
            view = dbobject.GetSummaryInformation(200)
            return view.GetProperty(n)

    print("Reading msi is completed....!")
    path = os.path.join(os.getcwd(), MSI)
    Upgrde = GetMsiProperty(path, "UpgradeCode")
    ProductCode = GetMsiProperty(path, "ProductCode")
    UPN = Summaryinfo(path, 2)
    Application_Name = application[1]
    vendor_Name = application[5]
    NameSME=rtf[15]
    EmailSME=rtf[17]
    ContSME=rtf[19]

    if ("64-" in MSI):
        Bi = 64
    else:
        Bi = 32
    BU = UPN[len(UPN) - 6:len(UPN) - 3]

    def main():
        template_file_path = os.path.join(os.getcwd(), "Shiva.docx")
        output_file_path = os.path.join(os.getcwd(), str(UPN.decode("utf-8")) + "_PackageReport.docx")

        variables = {
            # "${EMPLOEE_NAME}": input("Enter your name :"),
            "${UPN}": str(UPN.decode("utf-8")),
            "${Ugrade code}": Upgrde,
            "${Product code}": ProductCode,
            "${Date}": today,
            "${MSI1}": "Msiexec.exe /i " + str(UPN.decode("utf-8")) + ".msi"+" /qb",
            "${Uninstall}": Ugrades,
            "${Bi}": str(Bi),
            "${BU}": str(BU.decode("utf-8")),
            "${Location}": str(Location),
            "${Method}": "Setup Capture",
            "${PRE}": pre,
            "${COM}": comp,
            "${A_Name}": str(Application_Name),
            "${V_Name}": vendor_Name,
            "${NameSME}":NameSME,
            "${EmailSME}":EmailSME,
            "${ContSME}":ContSME,
            "${Scr}": "N/A"
            # "${MB}":ss


        }

        print("Entering to print variable in package report....!")
        template_document = docx.Document(template_file_path)

        for variable_key, variable_value in variables.items():
            for paragraph in template_document.paragraphs:
                paragraph.text = paragraph.text.replace(variable_key, variable_value)

            for table in template_document.tables:
                for col in table.columns:
                    for cell in col.cells:
                        for paragraph in cell.paragraphs:
                            if variable_key in paragraph.text:
                                paragraph.text = paragraph.text.replace(variable_key, variable_value)
                            else:
                                pass
        print("Header is entering...!")
        section = template_document.sections[0]
        header = section.header
        header_para = header.paragraphs[0]
        header_para.text = str(UPN.decode("utf-8"))+"Global Application Packaging Report Template 1.3.1.doc"

        line = ['Statistics', UPN, 'Manual Installation Instructions', 'Automated Install Command Line',
                'Automated Uninstall Command Line', 'Data Sources','Abbreviations' ,'Data Sources', 'File Access',
                'Middleware Configuration']

        li = ['Package Build Report v1.A for','Abbreviations', 'Summary Details', 'Change History', 'Application Details',
             'Package Details','Prerequisites and Dependencies', 'Security', 'Uninstalls/Upgrade', 'Scripting Additions',
             'Package Installation','MSI Logging', 'Verification Tests', 'Known Issues & Restrictions ',
             'Client/Server Application - Data Connections','References ', 'Analysis Aids']

        for i in li:
            for paragraph in template_document.paragraphs:
                if i == paragraph.text:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.size = Pt(16)
                        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                    # paragraph.add_run().bold =True
                    # print(paragraph.text)
        for i in line:
            for paragraph in template_document.paragraphs:
                if i == paragraph.text:
                    for run in paragraph.runs:
                        run.bold = True
                        if run.text == UPN:
                            print("This", run.text)
                        else:
                            run.underline = True
        print("Created format...!")
        template_document.save(output_file_path)
        print('Created package report for MSI....!')

    if __name__ == '__main__':
        main()

elif len(l) == 1 and l[0].endswith(".mst"):
    MST = l[0]
    if (MST[:4] == "ZIC-"):
        Location = "172.16.64.2\Pkg-I\Zurich\pkg"
    else:
        Location = "172.16.64.2\Pkg-I\Zurich\pkg-citrix"
    if MST.endswith(".mst"):
        for file in os.listdir():
            if file.endswith(".msi"):
                MSI = file

                def GetMsiProperty(path, property):
                    db = OpenDatabase(path, MSIDBOPEN_READONLY)
                    view = db.OpenView("SELECT Value FROM Property WHERE Property='" + property + "'")
                    view.Execute(None)
                    result = view.Fetch()
                    # print dir(result)
                    return result.GetString(1)

                path = os.path.join(os.getcwd(), MSI)
                Upgrde = GetMsiProperty(path, "UpgradeCode")
                ProductCode = GetMsiProperty(path, "ProductCode")
                Application_Name = application[1]
                vendor_Name = application[5]
                NameSME = rtf[15]
                EmailSME = rtf[17]
                ContSME = rtf[19]

                n = len(MST)
                UPN = MST[0:n - 4]
                BU = UPN[len(UPN) - 6:len(UPN) - 3]
                if "64-" in MST:
                    Bi = 64
                else:
                    Bi = 32

                """
                size = 0
                Folderpath = "\\" + "\\" + Location + "\\" + UPN
                for path, dirs, files in os.walk(Folderpath):
                    for f in files:
                        fp = os.path.join(path, f)
                        size += os.stat(fp).st_size
                        if size == 0:
                            ss == 0
                        else:
                            size_name = ("B", "KB", "MB", "GB", "TB", "PB", "EB", "ZB", "YB")
                            i = int(math.floor(math.log(size, 1024)))
                            p = math.pow(1024, i)
                            s = round(size / p, 2)
                            ss = ("%s %s" % (s, size_name[i]))
"""

                def main():
                    template_file_path = os.path.join(os.getcwd(), "Shiva.docx")
                    output_file_path = os.path.join(os.getcwd(), UPN + "_PackageReport.docx")
                    variables = {
                        # "${EMPLOEE_NAME}": input("Enter your name :"),
                        "${Bi}": str(Bi),
                        "${BU}": str(BU),
                        "${UPN}": str(UPN),
                        "${Ugrade code}": Upgrde,
                        "${Product code}": ProductCode,
                        "${Date}": str(today),
                        "${MSI1}": "Msiexec.exe /i "+ MSI +"TRANSFORMS=" +str(MST)+" /qb",
                        "${Location}": str(Location),
                        "${Method}": "Mst",
                        "${PRE}": pre,
                        "${COM}": comp,
                        "${A_Name}": str(Application_Name),
                        "${V_Name}": vendor_Name,
                        "${Uninstall}": Ugrades,
                        "${NameSME}": NameSME,
                        "${EmailSME}": EmailSME,
                        "${ContSME}": ContSME,
                        "${Scr}": "N/A"
                        # "${MB}": ss
                    }

                    template_document = docx.Document(template_file_path)

                    for variable_key, variable_value in variables.items():
                        for paragraph in template_document.paragraphs:
                            paragraph.text = paragraph.text.replace(variable_key, variable_value)
                            # replace_text_in_paragraph(paragraph, variable_key, variable_value)

                        for table in template_document.tables:
                            for col in table.columns:
                                for cell in col.cells:
                                    for paragraph in cell.paragraphs:
                                        if variable_key in paragraph.text:
                                            paragraph.text = paragraph.text.replace(variable_key, variable_value)
                                        else:
                                            pass
                                        # replace_text_in_paragraph(paragraph, variable_key, variable_value)
                    section = template_document.sections[0]
                    header = section.header
                    header_para = header.paragraphs[0]
                    header_para.text = UPN + "_Global Application Packaging Report Template 1.3.1.doc"

                    line = ['Statistics', UPN, 'Manual Installation Instructions', 'Automated Install Command Line',
                            'Automated Uninstall Command Line', 'Data Sources', 'Abbreviations', 'Data Sources',
                            'File Access',
                            'Middleware Configuration']

                    li = ['Package Build Report v1.A for','Abbreviations', 'Summary Details', 'Change History', 'Application Details',
                          'Package Details','Prerequisites and Dependencies', 'Security', 'Uninstalls/Upgrade', 'Scripting Additions',
                          'Package Installation','MSI Logging', 'Verification Tests', 'Known Issues & Restrictions ',
                          'Client/Server Application - Data Connections','References ', 'Analysis Aids']
                    for i in li:
                        for paragraph in template_document.paragraphs:
                            if i == paragraph.text:
                                for run in paragraph.runs:
                                    run.bold = True
                                    run.font.size = Pt(16)
                                    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                                # paragraph.add_run().bold =True
                                # print(paragraph.text)
                    for i in line:
                        for paragraph in template_document.paragraphs:
                            if i == paragraph.text:
                                for run in paragraph.runs:
                                    run.bold = True
                                    if run.text == UPN:
                                        print("This", run.text)
                                    else:
                                        run.underline = True
                    print("Created format...!")

                    template_document.save(output_file_path)
                    print('Created package report for MST')


                if __name__ == '__main__':
                    main()

elif len(l) == 1 and l[0].endswith(".exe"):
    EXE=l[0]
    if (EXE[:4] == "ZIC-"):
        Location = "172.16.64.2\Pkg-I\Zurich\pkg"
    else:
        Location = "172.16.64.2\Pkg-I\Zurich\pkg-citrix"
    for file in os.listdir():
        if (file.startswith("ZIC-") or file.startswith("ZCX-")) and (file.endswith(".docx")):
            doc=file
            n=len(doc)
            upn=(str(doc[0:n - 18]))
    Application_Name = application[1]
    vendor_Name = application[5]
    UPN=upn
    NameSME = rtf[15]
    EmailSME = rtf[17]
    ContSME = rtf[19]
    BU = UPN[len(UPN) - 6:len(UPN) - 3]
    if "64-" in UPN:
        Bi = 64
    else:
        Bi = 32
    """size = 0
    Folderpath = "\\" + "\\" + Location + "\\" + UPN
    for path, dirs, files in os.walk(Folderpath):
        for f in files:
            fp = os.path.join(path, f)
            size += os.stat(fp).st_size
            if size == 0:
                ss == 0
            else:
                size_name = ("B", "KB", "MB", "GB", "TB", "PB", "EB", "ZB", "YB")
                i = int(math.floor(math.log(size, 1024)))
                p = math.pow(1024, i)
                s = round(size / p, 2)
                ss = ("%s %s" % (s, size_name[i]))
"""

    def main():
        template_file_path = os.path.join(os.getcwd(), "Shiva.docx")
        output_file_path = os.path.join(os.getcwd(), UPN + "_PackageReport.docx")
        variables = {
            # "${EMPLOEE_NAME}": input("Enter your name :"),
            "${Bi}": str(Bi),
            "${BU}": str(BU),
            "${UPN}": str(UPN),
            "${Ugrade code}": "N/A",
            "${Product code}": "N/A",
            "${Date}": str(today),
            "${MSI1}": EXE +" /s",
            "${Location}": str(Location),
            "${Method}": "Wrapper",
            "${PRE}": pre,
            "${COM}": comp,
            "${A_Name}": str(Application_Name),
            "${V_Name}": vendor_Name,
            "${Uninstall}": Ugrades,
            "${NameSME}": NameSME,
            "${EmailSME}": EmailSME,
            "${ContSME}": ContSME,
            "${Scr}": UPN + "_Install.wse – used to install the package" + "\n" + UPN + "_Uninstall.wse – used to uninstall the package."
            # "${MB}": ss
        }

        template_document = docx.Document(template_file_path)

        for variable_key, variable_value in variables.items():
            for paragraph in template_document.paragraphs:
                paragraph.text = paragraph.text.replace(variable_key, variable_value)
                # replace_text_in_paragraph(paragraph, variable_key, variable_value)

            for table in template_document.tables:
                for col in table.columns:
                    for cell in col.cells:
                        for paragraph in cell.paragraphs:
                            if variable_key in paragraph.text:
                                paragraph.text = paragraph.text.replace(variable_key, variable_value)
                            else:
                                pass
                            # replace_text_in_paragraph(paragraph, variable_key, variable_value)
        section = template_document.sections[0]
        header = section.header
        header_para = header.paragraphs[0]
        header_para.text = UPN + "_Global Application Packaging Report Template 1.3.1.doc"
        # print(header_para.text)

        line = ['Statistics', UPN, 'Manual Installation Instructions', 'Automated Install Command Line',
                'Automated Uninstall Command Line', 'Data Sources', 'Abbreviations', 'Data Sources',
                'File Access',
                'Middleware Configuration']

        li = ['Package Build Report v1.A for','Abbreviations', 'Summary Details', 'Change History', 'Application Details',
              'Package Details','Prerequisites and Dependencies', 'Security', 'Uninstalls/Upgrade', 'Scripting Additions',
              'Package Installation','MSI Logging', 'Verification Tests', 'Known Issues & Restrictions ',
              'Client/Server Application - Data Connections','References ', 'Analysis Aids']

        for i in li:
            for paragraph in template_document.paragraphs:
                if i == paragraph.text:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.size = Pt(16)
                        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                    # paragraph.add_run().bold =True
                    # print(paragraph.text)
        for i in line:
            for paragraph in template_document.paragraphs:
                if i == paragraph.text:
                    for run in paragraph.runs:
                        run.bold = True
                        if run.text == UPN:
                            print("This", run.text)
                        else:
                            run.underline = True
        print("Created format...!")

        template_document.save(output_file_path)
        print('Created package report for EXE...!')


    if __name__ == '__main__':
        main()



